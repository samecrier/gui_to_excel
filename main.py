import tkinter as tk
import tkinter.scrolledtext
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from docx.shared import Cm, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from python_docx_replace import docx_replace
from collections import defaultdict

import datetime
import openpyxl as op
import pyperclip as clip
from openpyxl.cell.cell import Cell
from openpyxl.styles import Border, Side, Alignment, Font
import csv
import docx
import re
from ctypes import windll  # установка разрешения
import sys
import os

path_1 = ''
path_2 = ''
sleep_x = 0


@staticmethod
def keypress(event):
	if event.keycode == 86:
		event.widget.event_generate('<<Paste>>')
	elif event.keycode == 88:
		event.widget.event_generate('<<Cut>>')
	elif event.keycode == 65:
		event.widget.event_generate('<<SelectAll>>')


def read_csv_one_string(filename, delimiter=';'):
	with open(filename, 'r', encoding='utf-8', newline='') as f:
		csv_reader = csv.reader(f, delimiter=delimiter)
		for row in csv_reader:
			f.close()
			return row


def read_csv_full(filename, delimiter='&'):
	with open(filename, 'r', encoding='utf-8', newline='') as f:
		csv_reader = csv.reader(f, delimiter=delimiter)
		csv_list = []
		for row in csv_reader:
			csv_list.append(row)
		f.close()
		return csv_list


def write_csv(row, filename, delimiter=';'):
	row = sorted(row)
	with open(filename, 'w', newline='', encoding='utf-8') as f:
		writer = csv.writer(f, delimiter=delimiter)
		writer.writerow(row)
		f.close()


def write_settings_csv(value, row_number, filename='datas/settings.csv'):
	settings_list = read_csv_one_string(filename='datas/settings.csv')
	settings_list[row_number] = value
	with open(filename, 'w', newline='', encoding='utf-8') as f:
		writer = csv.writer(f, delimiter=';')
		writer.writerow(settings_list)
		f.close()


def write_dict_to_list(dict_for_func, filename='datas/indicators_to_code.csv', type_record='w'):
	dict_to_list = [[k, v] for k, v in dict_for_func.items()]
	with open(file=filename, mode=type_record, newline='', encoding='utf-8') as f:
		for row in dict_to_list:
			writer = csv.writer(f, delimiter='&')
			writer.writerow(row)
		f.close()


def read_csv_to_dict(filename='datas/indicators_to_code.csv'):
	csv_to_dict = {}
	with open(file=filename, newline='', encoding='utf-8') as f:
		dict_to_list = csv.reader(f, delimiter='&')
		for row in dict_to_list:
			csv_to_dict[row[0]] = row[1]
		f.close()
	return csv_to_dict


try:
	if read_csv_one_string(filename='datas/settings.csv')[1] == '':
		base_path1 = 'docs/Журнал_пробоподготовки,_исследования_проб_образцов_и_регистрации.xlsx'
	else:
		base_path1 = read_csv_one_string(filename='datas/settings.csv')[1]
except FileNotFoundError:
	with open('datas/settings.csv', 'w', newline='', encoding='utf-8') as f:
		writer = csv.writer(f, delimiter=';')
		writer.writerow([
			1.0,
			'',
			'',
			'Кулемин И.А.'
		])
	if read_csv_one_string(filename='datas/settings.csv')[1] == '':
		base_path1 = 'docs/Журнал_пробоподготовки,_исследования_проб_образцов_и_регистрации.xlsx'
	else:
		base_path1 = read_csv_one_string(filename='datas/settings.csv')[1]
	if read_csv_one_string(filename='datas/settings.csv')[2] == '':
		base_path2 = 'docs/Журнал_регистрации_проб_паразитологической_лаборатории_2023.xlsx'
	else:
		base_path2 = read_csv_one_string(filename='datas/settings.csv')[2]
else:
	if read_csv_one_string(filename='datas/settings.csv')[2] == '':
		base_path2 = 'docs/Журнал_регистрации_проб_паразитологической_лаборатории_2023.xlsx'
	else:
		base_path2 = read_csv_one_string(filename='datas/settings.csv')[2]


def settings_window():
	def bases_file_1():
		global base_path1
		ebase1_path['state'] = tk.NORMAL
		base_path1 = filedialog.askopenfilename()
		ask_or = messagebox.askokcancel('Инфо', 'Вы точно хотите сменить файл пробоподготовки?',
		                                parent=window_for_settings)
		if ask_or:
			write_settings_csv(base_path1, row_number=1)
			ebase1_path.delete(0, tk.END)
			ebase1_path.insert(0, base_path1)
			ebase1_path['state'] = tk.DISABLED

	def bases_file_2():
		global base_path2
		ebase2_path['state'] = tk.NORMAL
		base_path2 = filedialog.askopenfilename()
		ask_or = messagebox.askokcancel('Инфо', 'Вы точно хотите сменить регистрационный файл?',
		                                parent=window_for_settings)
		if ask_or:
			write_settings_csv(base_path2, row_number=2)
			ebase2_path.delete(0, tk.END)
			ebase2_path.insert(0, base_path2)
			ebase2_path['state'] = tk.DISABLED

	def change_head_of_department():
		global head_of_department
		ask_or = messagebox.askokcancel('Инфо', 'Вы точно хотите сменить регистрационный файл?',
		                                parent=window_for_settings)
		if ask_or:
			write_settings_csv(head_of_label_entry.get(), row_number=3)
			head_of_department = read_csv_one_string(filename='datas/settings.csv')[3]
			head_of_label.config(text=head_of_department)
			head_of_label_entry.delete(0, tk.END)

	def scaling_option(scaling_number):
		win.tk.call('tk', 'scaling', scaling_number)
		write_settings_csv(value=scaling_number, row_number=0)
		check_scaling_messagebox = messagebox.askokcancel('Предупреждение',
		                                                  'Чтобы применение вступило в силу необходимо перезагрузить приложение. Нажмите ОК если хотите перезагрузить приложение сейчас (введенные данные будут утеряны). Нажмите Отмена, если хотите самостоятельно перезагрузить приложение.',
		                                                  parent=window_for_settings)
		if check_scaling_messagebox:
			python = sys.executable
			os.execl(python, python, *sys.argv)

	def indicator_to_det_nd():
		list_of_value = []

		def choose_det_nd(evt):
			try:
				if dict_data_set[list_of_value[-1]] != textbox_det_nd.get('1.0', 'end-1c'):
					ask_or = messagebox.askokcancel('Действие', 'Принять изменения или отменить?',
					                                parent=window_for_code_to_nd)
					if ask_or:
						dict_data_set[list_of_value[-1]] = textbox_det_nd.get('1.0', 'end-1c')
						write_dict_to_list(dict_for_func=dict_data_set)
			except (IndexError, KeyError):
				pass
			textbox_det_nd.delete(0.0, tk.END)
			w = evt.widget
			value = w.get(int(w.curselection()[0]))
			try:
				if list_of_value[-1] != value:
					list_of_value.append(value)
			except IndexError:
				list_of_value.append(value)
			textbox_det_nd.insert(tk.INSERT, dict_data_set[value])
			b1_delete_indicator.grid(row=1, column=8, columnspan=4, stick='w')

		def add_indicator():
			new_indicator = b0_entry_add_indicator.get()
			if new_indicator not in list(dict_data_set.keys()):
				dict_data_set[new_indicator] = ''
				indicators_list.insert(0, new_indicator)
				b0_entry_add_indicator.delete(0, tk.END)
				write_dict_to_list(dict_for_func=dict_data_set)
			else:
				messagebox.showerror('Ошибка', 'Этот индикатор уже есть в словаре', parent=window_for_code_to_nd)

		def del_indicator():
			selection = indicators_list.curselection()
			value = indicators_list.get(int(selection[0]))
			del dict_data_set[value]
			indicators_list.delete(selection[0])
			write_dict_to_list(dict_for_func=dict_data_set)

		def confirm_code_det_nd_changes():
			selection = indicators_list.curselection()
			value = indicators_list.get(int(indicators_list.curselection()[0]))
			dict_data_set[value] = textbox_det_nd.get('1.0', 'end-1c')
			write_dict_to_list(dict_for_func=dict_data_set)

		def close_window():
			try:
				if dict_data_set[list_of_value[-1]] != textbox_det_nd.get('1.0', 'end-1c'):
					ask_or = messagebox.askokcancel('Действие',
					                                'Обнаружено несохраненное изменение в последней записи. Принять ее изменения?',
					                                parent=window_for_code_to_nd)
					if ask_or:
						dict_data_set[list_of_value[-1]] = textbox_det_nd.get('1.0', 'end-1c')
						write_dict_to_list(dict_for_func=dict_data_set)
			except IndexError:
				pass
			except KeyError:
				write_dict_to_list(dict_for_func=dict_data_set)
			window_for_code_to_nd.destroy()

		window_for_code_to_nd = tk.Toplevel(window_for_settings)  # нельзя нажимать в других окнах
		window_for_code_to_nd.title('Словарик кодов')
		window_for_code_to_nd.geometry(f'{int(875 * scaling)}x{int(325.0 * scaling)}+1000+350')
		window_for_code_to_nd.protocol('WM_DELETE_WINDOW', close_window)  # закрытие приложения

		window_for_code_to_nd.bind("<Control-KeyPress>", keypress)
		window_for_code_to_nd.bind_all("<Control-KeyPress>", _copy)

		dict_data_set = {k[0]: k[1] for k in read_csv_full('datas/indicators_to_code.csv', delimiter='&')}
		indicator_list_var = tk.Variable(value=sorted(dict_data_set.keys()))
		indicators_list = tk.Listbox(window_for_code_to_nd, listvariable=indicator_list_var, exportselection=False)
		indicators_list.grid(row=0, column=0, columnspan=12, padx=42)
		indicators_list['width'] = 75  # надо будет по максимально фильтровать
		indicators_list.bind('<<ListboxSelect>>', choose_det_nd)

		textbox_det_nd = tk.Text(window_for_code_to_nd, wrap=tk.WORD, width=75, height=26)
		textbox_det_nd.grid(row=0, column=12)

		b0_entry_add_indicator = tk.Entry(window_for_code_to_nd, width=50, justify=tk.LEFT)
		b0_entry_add_indicator.grid(row=1, column=0, columnspan=4, stick='w', padx=0)

		b0_add_indicator = tk.Button(window_for_code_to_nd, text='Добавить показатель', command=add_indicator)
		b0_add_indicator.grid(row=1, column=4, columnspan=4, stick='w', padx=0)

		b1_delete_indicator = tk.Button(window_for_code_to_nd, text='Удалить показатель', command=del_indicator)

		b2_show_txt = tk.Button(window_for_code_to_nd, text='Принять изменение', font=('Arial', '12'),
		                        command=confirm_code_det_nd_changes)
		b2_show_txt.grid(row=1, column=12, stick='e')

	window_for_settings = tk.Toplevel(win)  # нельзя нажимать в других окнах
	window_for_settings.title('Настройки')
	window_for_settings.geometry(f'{int(430 * scaling)}x{int(325.0 * scaling)}+1000+350')
	window_for_settings.protocol('WM_DELETE_WINDOW')  # закрытие приложения

	tk.Label(window_for_settings, text='Введите коэффициент масштабирования').grid(row=0, column=0)
	tk.Label(window_for_settings, text=f'Текущий значение - {scaling}').grid(row=1, column=0)
	scaling_entry = tk.Entry(window_for_settings)
	scaling_entry.grid(row=2, column=0)
	tk.Button(window_for_settings, text='Применить', command=lambda: scaling_option(float(scaling_entry.get()))).grid(
		row=3, column=0)

	tk.Label(window_for_settings, text='Выбери файл прободготовки для выгрузки в базу').grid(row=4, column=0)
	ebase1_path = tk.Entry(window_for_settings, justify=tk.CENTER, font=('Arial', 10), width=80)
	ebase1_path.insert(0, base_path1)
	ebase1_path['state'] = tk.DISABLED
	ebase1_path.grid(row=5, column=0)
	tk.Button(window_for_settings, text='Выберите файл', bd=5, font=('Arial', 10), command=bases_file_1).grid(row=6,
	                                                                                                          column=0)

	tk.Label(window_for_settings, text='Выбери регистрационный файл для выгрузки в базу').grid(row=7, column=0)
	ebase2_path = tk.Entry(window_for_settings, justify=tk.CENTER, font=('Arial', 10), width=80)
	ebase2_path.insert(0, base_path2)
	ebase2_path['state'] = tk.DISABLED
	ebase2_path.grid(row=8, column=0)
	tk.Button(window_for_settings, text='Выберите файл', bd=5, font=('Arial', 10), command=bases_file_2).grid(row=9,
	                                                                                                          column=0)

	tk.Label(window_for_settings, text='Заведующий паразитологической лабораторией').grid(row=10, column=0)
	head_of_label = tk.Label(window_for_settings, text=head_of_department)
	head_of_label.grid(row=11, column=0)
	head_of_label_entry = tk.Entry(window_for_settings)
	head_of_label_entry.grid(row=12, column=0)
	tk.Button(window_for_settings, text='Применить', command=change_head_of_department).grid(row=13, column=0)

	tk.Button(window_for_settings, text='Открыть словарик код-показатель', font=('Arial', '12'),
	          command=indicator_to_det_nd).grid(row=14, column=0, pady=20)


scaling = float(read_csv_one_string('datas/settings.csv')[0])
head_of_department = read_csv_one_string('datas/settings.csv')[3]

win = tk.Tk()
text_copy = tkinter.scrolledtext.ScrolledText(master=win, wrap='none')


def _copy(event):
	if event.keycode == 67:
		try:
			string = text_copy.selection_get()
			clip.copy(string)
		except:
			pass


windll.shcore.SetProcessDpiAwareness(1)
win.tk.call('tk', 'scaling', scaling)
# отменить скейлинг
win.geometry(f'{int(800.0 * scaling)}x{int(550.0 * scaling)}+50+50')
win.title('Программа')
win.event_delete('<<Paste>>', '<Control-V>')
win.event_delete('<<Copy>>', '<Control-C>')
win.event_delete('<<Cut>>', '<Control-X>')
win.event_delete('<<SelectAll>>', '<Control-A>')
win.event_delete('<<Paste>>', '<Control-v>')
win.event_delete('<<Copy>>', '<Control-c>')
win.event_delete('<<Cut>>', '<Control-x>')
win.event_delete('<<SelectAll>>', '<Control-a>')
win.bind("<Control-KeyPress>", keypress)
win.bind_all("<Control-KeyPress>", _copy)


def get_info():
	print(f'Номер лабораторного журнала - {nb_lab_journal.get()}')
	print(f'Регистрационный номер пробы - {rg_nb_sample.get()}')
	print(f'Наименование пробы(образца) - {name_sample.get()}')
	print(f'ФИО специалиста ответственного за пробоподготовку - {nm_sample_executor.get()}')
	print(f'Примечания пробоподготовки - {nt_sample.get()}')
	print(f'Примечания регистрационного журнала - {nt_register.get()}')
	print(f'Перечень показателей через запятую - {ls_indicators.get()}')
	print(f'Реквизиты НД для проведения пробоподготовки - {det_nd_prep_sample.get()}')
	print(f'Реквизиты НД на метод исследования - {det_nd_research.get()}')
	print(f'ФИО специалиста проводившего исследование - {sp_did_research.get()}')
	print(f'ФИО ответственного исполнителя - {rsp_executor.get()}')
	print(f'Дата начала исследования - {dt_st_research.get()}')
	print(f'Дата начала пробоподготовки - {dt_st_sample_prep.get()}')
	print(f'Дата отбора пробы (образца) - {dt_st_sampling.get()}')
	print(f'Дата поступления - {dt_get_receipt.get()}')
	print(f'Дата окончания исследования - {dt_fn_research.get()}')
	print(f'Дата окончания пробоподготовки - {dt_fn_sample_prep.get()}')
	print(f'Дата утилизации пробы/сведения о консервации - {dt_disposal.get()}')
	print(f'Дата выписки листа протокола - {dt_issue_protocol.get()}')
	print(f'Этапы пробоподготовки - {steps_sample.get()}')
	print(f'Этапы исследования - {stp_research.get()}')
	print('_________________________________________________')


def write_history(new_csv, type_data='list', type_record='a', filename='datas/query_history.csv'):
	with open(file=filename, mode=type_record, newline='', encoding='utf-8') as f:
		if type_data == 'list':
			for row in new_csv:
				writer = csv.writer(f, delimiter='&')
				writer.writerow(row)
		if type_data == 'row':
			writer = csv.writer(f, delimiter='&')
			writer.writerow(new_csv)
		f.close()


def excel_func():
	# try:
	def styled_cells(data, sheet):
		if len(data) == 15:
			for i, styled_cell in enumerate(data):
				if i == 0:
					styled_cell = int(styled_cell)
				if i in (5, 6, 10, 11):
					styled_cell = datetime.datetime(int(styled_cell.split('.')[-1]), int(styled_cell.split('.')[1]),
					                                int(styled_cell.split('.')[0]), 0, 0)
				styled_cell = Cell(sheet, column="A", value=styled_cell)
				styled_cell.font = Font(name='Calibri', size=11)
				styled_cell.border = Border(left=Side(style='thin'),
				                            right=Side(style='thin'),
				                            top=Side(style='thin'),
				                            bottom=Side(style='thin'))
				styled_cell.alignment = Alignment(vertical='bottom', wrap_text=True)
				if i in (5, 6, 10, 11):
					styled_cell.number_format = 'dd/mm/yyyy;@'
					styled_cell.alignment = Alignment(horizontal='right')
				yield styled_cell
		else:
			for i, styled_cell in enumerate(data):
				if i == 0:
					styled_cell = int(styled_cell)
				if i in (3, 4, 5, 7, 8, 9):
					styled_cell = datetime.datetime(int(styled_cell.split('.')[-1]), int(styled_cell.split('.')[1]),
					                                int(styled_cell.split('.')[0]), 0, 0)
				styled_cell = Cell(sheet, column="A", value=styled_cell)
				styled_cell.font = Font(name='Calibri', size=11)
				styled_cell.border = Border(left=Side(style='thin'),
				                            right=Side(style='thin'),
				                            top=Side(style='thin'),
				                            bottom=Side(style='thin'))
				styled_cell.alignment = Alignment(vertical='bottom', wrap_text=True)
				if i in (3, 4, 5, 7, 8, 9):
					styled_cell.number_format = 'dd/mm/yyyy;@'
					styled_cell.alignment = Alignment(horizontal='right')
				yield styled_cell

	if rg_nb_sample.get() == '':
		messagebox.showerror('Ошибка', 'Введите регистрационный номер пробы для отправки')
		return
	dict_keys = []
	with open('datas/query_history.csv', 'r', encoding='utf-8', newline='') as f:
		csv_reader = csv.reader(f, delimiter='&')
		for row in csv_reader:
			dict_keys.append(row[1])
		f.close()
	if rg_nb_sample.get() in dict_keys:
		answer = messagebox.askokcancel('Предупреждение',
		                                'Данный регистрационный номер уже находится в базе. Если вы хотите заменить запись нажмите ок, если вы не хотите заменять запись нажмите отмена',
		                                parent=win)
		if answer == True:
			old_csv = read_csv_full('datas/query_history.csv', delimiter='&')
			new_csv = []
			for row in old_csv:
				if rg_nb_sample.get() not in row:
					new_csv.append(row)
			write_history(new_csv, type_record='w')
		if answer == False:
			messagebox.showinfo('Информация', 'Изменение не будут применены к записи.')
			return

	if path_1 == '':
		path_sample_file = 'docs\\test_file_sample.xlsx'
	else:
		path_sample_file = path_1
	if path_2 == '':
		path_register_file = 'docs\\test_file_register.xlsx'
	else:
		path_register_file = path_2

	book_1 = op.load_workbook(filename=path_sample_file)
	sheet_1 = book_1.active
	book_2 = op.load_workbook(filename=path_register_file)
	sheet_2 = book_2.active

	if ('обнаружены' or 'не обнаружены') in ls_indicators.get():
		ls_indicators_research = ls_indicators.get()
	else:
		ls_indicators_research = ls_indicators.get() + ' ' + default_indicator

	sample_file = [
		nb_lab_journal.get(),  # 0
		rg_nb_sample.get(),  # 1
		name_sample.get(),  # 2
		det_nd_prep_sample.get(),  # 3
		steps_sample.get(),  # 4
		dt_st_sample_prep.get(),  # 5
		dt_fn_sample_prep.get(),  # 6
		nm_sample_executor.get(),  # 7
		det_nd_research.get(),  # 8
		stp_research.get(),  # 9
		dt_st_research.get(),  # 10
		dt_fn_research.get(),  # 11
		ls_indicators_research,  # 12
		sp_did_research.get(),  # 13
		nt_sample.get()  # 14
	]
	register_file = [
		nb_lab_journal.get(),  # 0
		rg_nb_sample.get(),  # 1
		name_sample.get(),  # 2
		dt_st_sampling.get(),  # 3
		dt_get_receipt.get(),  # 4
		dt_st_research.get(),  # 5
		ls_indicators.get(),  # 6
		dt_fn_research.get(),  # 7
		dt_disposal.get(),  # 8
		dt_issue_protocol.get(),  # 9
		rsp_executor.get(),  # 10
		nt_register.get()  # 11
	]

	sheet_1.append(styled_cells(sample_file, sheet_1))
	book_1.save(filename=path_sample_file)
	path = os.path.realpath(path_sample_file)

	sheet_2.append(styled_cells(register_file, sheet_2))
	book_2.save(filename=path_register_file)
	path = os.path.realpath(path_register_file)

	write_history(
		[
			nb_lab_journal.get(),  # 0
			rg_nb_sample.get(),  # 1
			name_sample.get(),  # 2
			nm_sample_executor.get(),  # 3
			nt_sample.get(),  # 4
			nt_register.get(),  # 5
			ls_indicators_research,  # 6
			det_nd_prep_sample.get(),  # 7
			det_nd_research.get(),  # 8
			sp_did_research.get(),  # 9
			rsp_executor.get(),  # 10
			dt_st_research.get(),  # 11
			dt_st_sample_prep.get(),  # 12
			dt_st_sampling.get(),  # 13
			dt_get_receipt.get(),  # 14
			dt_fn_research.get(),  # 15
			dt_fn_sample_prep.get(),  # 16
			dt_disposal.get(),  # 17
			dt_issue_protocol.get(),  # 18
			steps_sample.get(),  # 19
			stp_research.get(),  # 20
		], type_data='row'
	)

	if op_xl_button_value.get() == 'No':
		print('Сохранение в эксель без открытия')
		return
	os.startfile(path_sample_file)
	os.startfile(path_register_file)
	print('Сохранение в эксель с открытием ')


# except:
# 	messagebox.showerror('Ошибка', 'Введен неправильный формат данных', parent=win)


def get_file_1():
	global path_1
	e1_path['state'] = tk.NORMAL
	path_1 = filedialog.askopenfilename()
	# tk.Label(win, text=path_1, anchor='w', width=10, height=1).grid(row=21, column=1, stick='w')
	e1_path.delete(0, tk.END)
	e1_path.insert(0, path_1)
	e1_path['state'] = tk.DISABLED


def get_file_2():
	global path_2
	e2_path['state'] = tk.NORMAL
	path_2 = filedialog.askopenfilename()
	# tk.Label(win, text=path_2, anchor='w', width=20, height=1).grid(row=22, column=1, stick='w')
	e2_path.delete(0, tk.END)
	e2_path.insert(0, path_2)
	e2_path['state'] = tk.DISABLED


def repeat_for_nd():
	det_nd_research.delete(0, tk.END)
	if repeat_for_nd_value.get() == 'Yes':
		det_nd_research.insert(0, det_nd_prep_sample.get())


def nd_check_button_off(evt=None):
	if det_nd_research.get() != det_nd_prep_sample.get():
		repeat_for_nd_value.set('No')


def repeat_for_dt_st_1():
	dt_st_sample_prep.delete(0, tk.END)
	if dt_st_value_1.get() == 'Yes':
		dt_st_sample_prep.insert(0, dt_st_research.get())


def dt_st_1_check_off(evt=None):
	if dt_st_sample_prep.get() != dt_st_research.get():
		dt_st_value_1.set('No')


def repeat_for_dt_st_2():
	dt_st_sampling.delete(0, tk.END)
	if dt_st_value_2.get() == 'Yes':
		dt_st_sampling.insert(0, dt_st_research.get())


def dt_st_2_check_off(evt=None):
	if dt_st_sampling.get() != dt_st_research.get():
		dt_st_value_2.set('No')


def repeat_for_dt_st_3():
	dt_get_receipt.delete(0, tk.END)
	if dt_st_value_3.get() == 'Yes':
		dt_get_receipt.insert(0, dt_st_research.get())


def dt_st_3_check_off(evt=None):
	if dt_get_receipt.get() != dt_st_research.get():
		dt_st_value_3.set('No')


def check_st_functions(evt):
	dt_st_1_check_off()
	dt_st_2_check_off()
	dt_st_3_check_off()


def repeat_for_dt_fn_1():
	dt_fn_sample_prep.delete(0, tk.END)
	if dt_fn_value_1.get() == 'Yes':
		dt_fn_sample_prep.insert(0, dt_fn_research.get())


def dt_fn_1_check_off(evt=None):
	if dt_fn_sample_prep.get() != dt_fn_research.get():
		dt_fn_value_1.set('No')


def repeat_for_dt_fn_2():
	dt_disposal.delete(0, tk.END)
	if dt_fn_value_2.get() == 'Yes':
		dt_disposal.insert(0, dt_fn_research.get())


def dt_fn_2_check_off(evt=None):
	if dt_disposal.get() != dt_fn_research.get():
		dt_fn_value_2.set('No')


def repeat_for_dt_fn_3():
	dt_issue_protocol.delete(0, tk.END)
	if dt_fn_value_3.get() == 'Yes':
		dt_issue_protocol.insert(0, dt_fn_research.get())


def dt_fn_3_check_off(evt=None):
	if dt_issue_protocol.get() != dt_fn_research.get():
		dt_fn_value_3.set('No')


def check_fn_functions(evt):
	dt_fn_1_check_off()
	dt_fn_2_check_off()
	dt_fn_3_check_off()


def repeat_for_stp():
	global glb_stp_research_check_name, glb_stp_number_of_research
	dict_for_end = {
		1: 'препарат', 2: 'препарата', 3: 'препарата', 4: 'препарата', 5: 'препаратов', 6: 'препаратов',
		7: 'препаратов', 8: 'препаратов', 9: 'препаратов', 10: 'препаратов', 11: 'препаратов', 12: 'препаратов',
		13: 'препаратов', 14: 'препаратов', 15: 'препаратов', 16: 'препаратов', 17: 'препаратов', 18: 'препаратов',
		19: 'препаратов', 20: 'препаратов', 21: 'препарат', 22: 'препарата', 23: 'препарата', 24: 'препарата',
		25: 'препаратов', 26: 'препаратов', 27: 'препаратов', 28: 'препаратов', 29: 'препаратов', 30: 'препаратов',
		31: 'препарат', 32: 'препарата', 33: 'препарата', 34: 'препарата', 35: 'препаратов', 36: 'препаратов',
		37: 'препаратов', 38: 'препаратов', 39: 'препаратов', 40: 'препаратов', 41: 'препарат', 42: 'препарата',
		43: 'препарата', 44: 'препарата', 45: 'препаратов', 46: 'препаратов', 47: 'препаратов', 48: 'препаратов',
		49: 'препаратов', 50: 'препаратов', 51: 'препарат', 52: 'препарата', 53: 'препарата', 54: 'препарата',
		55: 'препаратов', 56: 'препаратов', 57: 'препаратов', 58: 'препаратов', 59: 'препаратов', 60: 'препаратов',
		61: 'препарат', 62: 'препарата', 63: 'препарата', 64: 'препарата', 65: 'препаратов', 66: 'препаратов',
		67: 'препаратов', 68: 'препаратов', 69: 'препаратов', 70: 'препаратов', 71: 'препарат', 72: 'препарата',
		73: 'препарата', 74: 'препарата', 75: 'препаратов', 76: 'препаратов', 77: 'препаратов', 78: 'препаратов',
		79: 'препаратов', 80: 'препаратов', 81: 'препарат', 82: 'препарата', 83: 'препарата', 84: 'препарата',
		85: 'препаратов', 86: 'препаратов', 87: 'препаратов', 88: 'препаратов', 89: 'препаратов', 90: 'препаратов',
		91: 'препарат', 92: 'препарата', 93: 'препарата', 94: 'препарата', 95: 'препаратов', 96: 'препаратов',
		97: 'препаратов', 98: 'препаратов', 99: 'препаратов', 100: 'препаратов'
	}
	if repeat_for_stp_value.get() == 'Yes':
		stp_research_result = steps_sample.get().split('; ')[-1]
		stp_research_result_check_digit = stp_research_result.split(' ')[0]
		if stp_research_result_check_digit[-2:].isdigit():
			stp_research.delete(0, tk.END)
			digit_for_end = int(stp_research_result_check_digit[-2:])
			preparat_end = dict_for_end[digit_for_end]
			if preparat_end == 'препарат':
				stp_research_name = 'исследование выполнено; ' + f'{stp_research_result_check_digit}' + ' ' + preparat_end + ' исследован'
			else:
				stp_research_name = 'исследование выполнено; ' + f'{stp_research_result_check_digit}' + ' ' + preparat_end + ' исследованы'
			stp_research.insert(0, stp_research_name)
			glb_stp_research_check_name = stp_research_name
			glb_stp_number_of_research = stp_research_result_check_digit

		else:
			messagebox.showerror('Ошибка',
			                     'Неправильная форма этапов пробоподготовки, введите результаты иследования вручную')
			repeat_for_stp_value.set('No')


def for_stp_check_off(evt=None):
	if stp_research.get() != glb_stp_research_check_name:
		repeat_for_stp_value.set('No')


def check_stp_function(evt=None):
	if glb_stp_number_of_research == '':
		return
	stp_string = steps_sample.get().split('; ')[-1].split(' ')[0]
	if glb_stp_number_of_research != stp_string:
		repeat_for_stp_value.set('No')


def find_not_find(evt):
	global default_indicator
	default_indicator = evt.widget.get()


def start_window_0(variable, filename):
	def delete():
		selection = employee_listbox.curselection()
		name_of_selection = employee_listbox.get(int(employee_listbox.curselection()[0]))
		if name_of_selection == variable.get():
			variable.delete(0, tk.END)
		employees.remove(name_of_selection)
		write_csv(employees, filename)
		# мы можем получить удаляемый элемент по индексу
		# selected_language = employee_listbox.get(selection[0])
		employee_listbox.delete(selection[0])

	# добавление нового элемента
	def add():
		new_employee = employee_entry.get()
		list_employees = read_csv_one_string(filename)
		if list_employees:
			if new_employee not in list_employees:
				write_csv(list_employees + [new_employee], filename)
				employees.append(new_employee)
			else:
				messagebox.showerror('Ошибка', 'Такой сотрудник уже в списке!', parent=new_window_0)
				return
		else:
			write_csv([new_employee], filename)
		employee_listbox.insert(0, new_employee)

	def show_print(evt):
		w = evt.widget
		value = w.get(int(w.curselection()[0]))

	def add_to_enter_box():
		variable.delete(0, tk.END)
		selection = employee_listbox.curselection()
		name_of_selection = employee_listbox.get(int(employee_listbox.curselection()[0]))
		variable.insert(0, name_of_selection)
		new_window_0.destroy()

	new_window_0 = tk.Toplevel(win)
	new_window_0.grab_set()  # нельзя нажимать в других окнах
	new_window_0.title('Окно 1')
	new_window_0.geometry(f'{int(400.0 * scaling)}x{int(300.0 * scaling)}+1600+350')
	new_window_0.protocol('WM_DELETE_WINDOW')  # закрытие приложения
	new_window_0.wm_attributes("-topmost", 1)  # чтобы повешать поверх все окон, но работает и без
	# текстовое поле и кнопка для добавления в список
	employee_entry = ttk.Entry(new_window_0)
	employee_entry.grid(column=0, stick='e', row=0, padx=6, pady=6, sticky='ew')
	ttk.Button(new_window_0, text="Добавить специалиста", command=add).grid(column=1, row=0, padx=6, pady=6)
	employees = read_csv_one_string(filename)
	employees_var = tk.Variable(new_window_0, value=employees)
	employee_listbox = tk.Listbox(new_window_0, listvariable=employees_var)
	employee_listbox.grid(row=1, column=0, stick='e', columnspan=2, sticky='ew', padx=5, pady=5)

	ttk.Button(new_window_0, text="Применить", command=add_to_enter_box).grid(row=2, column=0, stick='e', padx=5,
	                                                                          pady=5)
	ttk.Button(new_window_0, text="Удалить", command=delete).grid(row=2, column=1, padx=5, pady=5)


def on_closing_0(this_window):
	if messagebox.askokcancel('Выход из приложения', 'Хотите ли вы выйти из приложения?'):
		this_window.destroy()


def dict_from_csv():
	csv_dict = {}
	with open('datas/query_history.csv', 'r', encoding='utf-8', newline='') as f:
		csv_reader = csv.reader(f, delimiter='&')
		for row in csv_reader:
			dict_key = row[1]
			dict_values = row
			csv_dict[dict_key] = dict_values
		f.close()
	return csv_dict


def word_func(dict_for_word, history_window_0):
	def add_entry_to_code(index):
		nonlocal index_entry
		index_entry += 1
		index = index_entry
		dict_of_entry_buttons[index_entry].append(tk.Entry(add_to_dict_window, width=50))
		dict_of_entry_buttons[index_entry].append(tk.Entry(add_to_dict_window, width=50))
		dict_of_entry_buttons[index_entry][0].grid(row=index, column=1, padx=5)
		dict_of_entry_buttons[index_entry][1].grid(row=index, column=2)

	def delete_entry_to_code(index):
		nonlocal index_entry
		dict_of_entry_buttons[index_entry][0].destroy()
		dict_of_entry_buttons[index_entry][1].destroy()
		del dict_of_entry_buttons[index_entry]
		index_entry -= 1

	def func_add_to_dict(entry_dict, nd_dict):
		previous_dict = read_csv_to_dict()
		for values in entry_dict.values():
			code = values[0].get()
			name = values[1].get()
			if code == '':
				messagebox.showerror('Ошибка', 'Поле кода пустое', parent=add_to_dict_window)
				return
			if code not in previous_dict:
				if name == '':
					messagebox.showerror('Ошибка',
					                     f'Код {code} отстутствует в базе данных и поля показателя пустое. Введите пожалуйста показатель.',
					                     parent=add_to_dict_window)
					return
				else:
					previous_dict[code] = name
					nd_dict[code] = name

			else:
				ask_or = messagebox.askokcancel('Предупреждение',
				                                f'Предыдущим значением для {code.upper()} являлось {previous_dict[code].upper()}, использовать сохраненный вариант? Отмена приведет к использованию  - нового варианта {name}',
				                                parent=add_to_dict_window)
				if ask_or == False:
					previous_dict[code] = name
					nd_dict[code] = name
				else:
					nd_dict[code] = previous_dict[code]
		write_dict_to_list(previous_dict)
		var_to_sleep.set(1)

	dict_first_item = next(iter(dict_for_word.values()))
	executor = dict_first_item[9]
	sample_name = dict_first_item[1]
	nd_details = list(set([dict_first_item[7], dict_first_item[8]]))

	symb_list = []
	for row in nd_details:
		row = re.sub(r'П\.', r'п.', row)
		var_last_symb = ''
		var_symb = ''
		counter = 0
		counter_2 = 0
		for index, _ in enumerate(row):
			var_last_symb = _
			if re.search(r'[^А-Я\s]', var_symb) and re.match(r'[А-Я]', _):
				symb_list.append(var_symb.strip())
				var_symb = _
			elif re.match(r'[А-Я\s]', _):
				if var_symb == '' and re.match(r'\s', _):
					continue
				elif re.match(r'[А-Я]', _) and var_symb == '':
					if counter_2 != 0:
						counter_2 = 0
						symb_list.append(var_symb.strip())
					var_symb = _
				else:
					if var_symb != '':
						var_symb += _
					else:
						var_symb = _
			else:
				counter_2 += 1
				counter += 1
				var_symb += _
			if index == len(row) - 1:
				symb_list.append(var_symb.strip())
	symb_list = [re.sub(r'^(.*),$', r'\1', row) for row in symb_list]
	symb_list = [re.sub(r'^(.*)\.$', r'\1', row) for row in symb_list]
	refactor_nd_codes = read_csv_to_dict()
	try:
		sample_name_code = sample_name.split('-')
		if len(sample_name_code) > 3:
			print(sample_name_code)
			sample_name_code = ('-').join(sample_name_code[:-1])
		else:
			sample_name_code = sample_name
	except:
		sample_name_code = sample_name

	indicator_names = dict_first_item[6]
	if ' не обнаружены' in indicator_names:
		indicator_names = indicator_names.replace(' не обнаружены', '')
	else:
		indicator_names = indicator_names.replace(' обнаружены', '')
	indicator_names = indicator_names.split(', ')

	nd_dict = {}

	#####АВТОМАТИЗИРОВАННЫЙ ВАРИАНТ ВЫБОРА КОДОВ - ПОКАЗАТЕЛЕЙ

	# def func_add_to_dict(name):
	# 	dict_data_set = {k[0]: k[1] for k in read_csv_full('datas/indicators_to_code.csv', delimiter='&')}
	# 	dict_data_set[name] = add_e0.get()
	# 	write_dict_to_list(dict_for_func=dict_data_set)
	# 	var_to_sleep.set(1)

	# for index, code in enumerate(list(set(symb_list))):
	# 	try:
	# 		nd_dict[code] = refactor_nd_codes[code]
	# 	except KeyError:
	# 		add_to_dict_window = tk.Toplevel(history_window_0)  # нельзя нажимать в других окнах
	# 		add_to_dict_window.title(f'Добавление {index+1} кода из {len(list(set(symb_list)))}')
	# 		add_to_dict_window.geometry(f'{int(550 * scaling)}x{int(325.0 * scaling)}+1000+350')
	# 		add_to_dict_window.protocol('WM_DELETE_WINDOW')
	# 		add_to_dict_window.bind("<Control-KeyPress>", keypress)
	# 		add_to_dict_window.bind_all("<Control-KeyPress>", _copy)
	#
	# 		tk.Label(add_to_dict_window,
	# 		         text=f'В базе данных для {code.upper()} не был обнаружен показатель. Введите показатель.').grid(row=0, column=0,
	# 		                                                                                           columnspan=50)
	# 		add_e0 = tk.Entry(add_to_dict_window, justify=tk.LEFT, width=50)
	# 		add_e0.grid(row=1, column=0, stick='w')
	# 		add_b0 = tk.Button(add_to_dict_window, text='Добавить', command=lambda: func_add_to_dict(code))
	# 		add_b0.grid(row=1, column=1, stick='w')
	# 		add_t0 = tk.Text(add_to_dict_window, width=100, wrap=tk.WORD)
	# 		add_t0.insert(tk.INSERT, 'Показатели из записи:\n\n')
	# 		add_t0.insert(tk.INSERT, dict_first_item[6])
	# 		add_t0['state'] = tk.DISABLED
	# 		add_t0.grid(row=2, column=0, columnspan=50, pady=10)
	# 		var_to_sleep = tk.IntVar()
	# 		var_to_sleep.set(0)
	# 		add_to_dict_window.wait_variable(var_to_sleep)
	# 		nd_dict[code] = add_e0.get()
	# 		add_to_dict_window.destroy()

	##########РУЧНОЙ ВАРИАНТ ВЫБОРА КОДОВ - ПОКАЗАТЕЛЕЙ
	add_to_dict_window = tk.Toplevel(history_window_0)  # нельзя нажимать в других окнах
	add_to_dict_window.title(f'Добавление показателей')
	add_to_dict_window.geometry(f'{int(850 * scaling)}x{int(325.0 * scaling)}+1000+350')
	add_to_dict_window.protocol('WM_DELETE_WINDOW')
	add_to_dict_window.bind("<Control-KeyPress>", keypress)
	add_to_dict_window.bind_all("<Control-KeyPress>", _copy)

	add_t0 = tk.Text(add_to_dict_window, width=50, wrap=tk.WORD)
	add_t0.insert(tk.INSERT, 'Показатели из записи:\n\n')
	add_t0.insert(tk.INSERT, dict_first_item[6] + '\n\n')
	add_t0.insert(tk.INSERT, 'Код из записи:\n\n')
	for row in symb_list:
		add_t0.insert(tk.INSERT, row + '\n')
	add_t0['state'] = tk.DISABLED
	add_t0.grid(row=0, rowspan=10, column=0, pady=10)
	var_to_sleep = tk.IntVar()
	var_to_sleep.set(0)

	dict_of_entry_buttons = defaultdict(list)
	dict_of_entry_buttons[2] = [tk.Entry(add_to_dict_window, width=50), tk.Entry(add_to_dict_window, width=50)]

	index_entry = 2
	add_b0 = tk.Button(add_to_dict_window, text='Добавить поле для кода-показателя',
	                   command=lambda: add_entry_to_code(index=index_entry))
	add_b0.grid(row=0, column=1, stick='nw')
	add_b0 = tk.Button(add_to_dict_window, text='Удалить поле для кода-показателя',
	                   command=lambda: delete_entry_to_code(index=index_entry))
	add_b0.grid(row=0, column=2, stick='nw')

	add_l0 = tk.Label(add_to_dict_window, text='Код', justify=tk.CENTER)
	add_l0.grid(row=1, column=1)
	add_l1 = tk.Label(add_to_dict_window, text='Показатель', justify=tk.CENTER)
	add_l1.grid(row=1, column=2)

	dict_of_entry_buttons[index_entry][0].grid(row=index_entry, column=1, padx=5)
	dict_of_entry_buttons[index_entry][1].grid(row=index_entry, column=2, padx=5)

	nd_dict = {}
	add_b1 = tk.Button(add_to_dict_window, text='Применить коды',
	                   command=lambda: func_add_to_dict(dict_of_entry_buttons, nd_dict))
	add_b1.grid(row=10, column=0)

	add_to_dict_window.wait_variable(var_to_sleep)

	###########################################
	indexes_nd_samples = len(dict_for_word)
	# list_samples = len(indicator_names)
	list_samples = len(nd_dict)
	print(indexes_nd_samples, list_samples)
	doc = docx.Document('docs/template.docx')

	def add_block():
		rows_for_table = indexes_nd_samples + (indexes_nd_samples * list_samples)
		table = doc.add_table(rows=1 + rows_for_table, cols=3, style="Table Grid")
		table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
		table.rows[0].height = Cm(1.4)

		for i, row in enumerate(table.rows):
			table.allow_autofit = False
			row.cells[0].width = Cm(5.19)
			row.cells[1].width = Cm(4.75)
			row.cells[2].width = Cm(6.5)

		def format_for_cell(variable, bold=False, align_p=False, align_v=False, pt=11):
			variable.paragraphs[0].runs[0].font.size = Pt(pt)
			if align_p != False:
				variable.paragraphs[0].alignment = align_v
			if align_v != False:
				variable.vertical_alignment = align_v
			if bold != False:
				variable.paragraphs[0].runs[0].font.bold = True

		def make_head():
			cell_0_0 = table.cell(0, 0)
			cell_0_0.text = 'Определяемые показатели'
			format_for_cell(cell_0_0, bold=True, align_p=WD_ALIGN_PARAGRAPH.CENTER, align_v=WD_ALIGN_VERTICAL.CENTER)

			cell_0_1 = table.cell(0, 1)
			cell_0_1.text = 'Результаты'
			format_for_cell(cell_0_1, bold=True, align_p=WD_ALIGN_PARAGRAPH.CENTER, align_v=WD_ALIGN_VERTICAL.CENTER)

			cell_0_2 = table.cell(0, 2)
			cell_0_2.text = 'НД на метод исследования'
			format_for_cell(cell_0_2, bold=True, align_p=WD_ALIGN_PARAGRAPH.CENTER, align_v=WD_ALIGN_VERTICAL.CENTER)

		def nd_samples_frame(i, fullname):
			row_1 = table.rows[i]
			cell_1_0, cell_1_1, cell_1_2 = row_1.cells[:3]
			cell_1_0.merge(cell_1_1)
			cell_1_0.merge(cell_1_2)
			cell_1_0 = table.cell(i, 0)
			cell_1_0.text = fullname
			format_for_cell(cell_1_0, bold=True)

		def ls_indicators_frame(i, nd_cell_code, nd_cell_result, nd_cell_name):
			table.rows[i].height = Cm(1.62)
			cell_2_0 = table.cell(i, 0)
			cell_2_0.text = nd_cell_name.capitalize()
			format_for_cell(cell_2_0, align_p=WD_ALIGN_PARAGRAPH.LEFT, align_v=WD_ALIGN_VERTICAL.CENTER)
			cell_2_1 = table.cell(i, 1)
			cell_2_1.text = nd_cell_result.capitalize()
			format_for_cell(cell_2_1, align_p=WD_ALIGN_PARAGRAPH.CENTER, align_v=WD_ALIGN_VERTICAL.CENTER)

			cell_2_2 = table.cell(i, 2)
			cell_2_2.text = nd_cell_code
			format_for_cell(cell_2_2, align_p=WD_ALIGN_PARAGRAPH.CENTER, align_v=WD_ALIGN_VERTICAL.CENTER)

		make_head()
		i = 1
		for key, value in dict_for_word.items():
			sample_fullname = f'{value[1]} / {value[2].capitalize()}'
			indicator_name = value[6]
			if ' не обнаружены' in indicator_name:
				indicator_result = 'Не обнаружено'
			else:
				indicator_result = 'Обнаружено'
			nd_samples_frame(i, fullname=sample_fullname)
			i += 1
			nd_dict_sorted = {k: v for k, v in sorted(nd_dict.items(), key=lambda item: item[1])}
			for dict_code, name in nd_dict_sorted.items():
				ls_indicators_frame(i, nd_cell_code=dict_code, nd_cell_name=name, nd_cell_result=indicator_result)
				i += 1

		doc.add_paragraph('')
		doc.add_paragraph('')

	add_block()

	# 2 таблица
	table_2 = doc.add_table(rows=2, cols=3, style="style_for_final")
	table_2.style.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

	for i, row in enumerate(table_2.rows):
		row.height = Cm(1.2)
		table_2.allow_autofit = False
		row.cells[0].width = Cm(7.5)
		row.cells[1].width = Cm(4.51)
		row.cells[2].width = Cm(4.49)

	cell2_0_0 = table_2.cell(0, 0)
	cell2_0_0.line_spacing = Pt(1.15)
	cell2_0_0.text = 'Уполномоченный специалист:\n'
	cell2_0_0.paragraphs[0].runs[0].font.size = Pt(12)
	cell2_0_0.paragraphs[0].runs[0].font.bold = True
	cell2_0_0.paragraphs[0].add_run('Врач-паразитолог')
	cell2_0_0.paragraphs[0].runs[1].font.size = Pt(11)
	cell2_0_0.paragraphs[0].runs[1].font.bold = False
	cell2_0_0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
	cell2_0_0.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

	cell2_0_2 = table_2.cell(0, 2)
	cell2_0_2.text = executor
	cell2_0_2.paragraphs[0].runs[0].font.size = Pt(11)
	cell2_0_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	cell2_0_2.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

	cell2_1_0 = table_2.cell(1, 0)
	cell2_1_0.text = 'Заведующий паразитологической лабораторией'
	cell2_1_0.paragraphs[0].runs[0].font.size = Pt(11)
	cell2_1_0.paragraphs[0].runs[0].font.bold = True
	cell2_1_0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
	cell2_1_0.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

	cell2_1_2 = table_2.cell(1, 2)
	cell2_1_2.text = head_of_department
	cell2_1_2.paragraphs[0].runs[0].font.size = Pt(11)
	cell2_1_2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	cell2_1_2.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

	protocol = doc.add_paragraph()
	protocol.text = '------------------------------------------------------------------конец протокола------------------------------------------------------------------'
	protocol.runs[0].font.size = Pt(9)

	docx_replace(doc, dt_st_sampling=dict_first_item[13], dt_st_research=dict_first_item[11],
	             dt_fn_research=dict_first_item[15])

	print('Word файл сгенерирован')
	doc.save(f'docs/{sample_name_code}.docx')
	messagebox.showinfo('Инфо', f'Word файл для {sample_name_code} сгенерирован', parent=add_to_dict_window)


def history_window():
	history_window_0 = tk.Toplevel(win)  # нельзя нажимать в других окнах
	history_window_0.title('Окно 1')
	history_window_0.geometry(f'{int(800.0 * scaling)}x{int(400.0 * scaling)}+1300+350')
	history_window_0.wm_attributes("-topmost", 0)  # чтобы повешать поверх все окон, но работает и без
	history_window_0.protocol('WM_DELETE_WINDOW')  # закрытие приложения

	def func_for_window():
		try:
			selection = l0.curselection()
			value = l0.get(int(l0.curselection()[0]))
		except:
			messagebox.showerror('Ошибка', 'Вы не выбрали группу для генерации', parent=history_window_0)
		else:
			dict_for_word = {k: v for k, v in history_dict.items() if re.fullmatch(f"{value}{r'-\d+'}", k)}
			l1 = dict_for_word.keys()
			d0 = {}
			for i in l1:
				row = i.split('-')
				d0[int(row[-1])] = i
			d0_new = dict(sorted(d0.items()))
			d0 = list(d0_new.values())
			dict_for_word = dict(sorted(dict_for_word.items(), key=lambda pair: d0.index(pair[0])))
			if len(dict_for_word) == 0:
				dict_for_word[value] = history_dict[value]
			word_func(dict_for_word, history_window_0=history_window_0)

	def confirm_to_main():
		global default_indicator
		selection = l1.curselection()
		value = l1.get(int(l1.curselection()[0]))
		for i in range(len(variables_for_row)):
			variables_for_row[i].delete(0, tk.END)
			if i == 6:
				indicator_names = history_dict[value][i]
				if ' не обнаружены' in indicator_names:
					indicator_names = indicator_names.replace(' не обнаружены', '')
					variables_for_row[i].insert(0, indicator_names)
					combo_indicators.current(0)
				else:
					indicator_names = indicator_names.replace(' обнаружены', '')
					variables_for_row[i].insert(0, indicator_names)
					combo_indicators.current(1)
			else:
				variables_for_row[i].insert(0, history_dict[value][i])

	def confirm_empty_to_main():
		selection = l1.curselection()
		value = l1.get(int(l1.curselection()[0]))
		for i in range(len(variables_for_row)):
			if variables_for_row[i].get() == '':
				if i == 6:
					indicator_names = history_dict[value][i]
					if ' не обнаружены' in indicator_names:
						indicator_names = indicator_names.replace(' не обнаружены', '')
						variables_for_row[i].insert(0, indicator_names)
						combo_indicators.current(0)
					else:
						indicator_names = indicator_names.replace(' обнаружены', '')
						variables_for_row[i].insert(0, indicator_names)
						combo_indicators.current(1)
				else:
					variables_for_row[i].insert(0, history_dict[value][i])

	def delete_from_csv():
		selection = l1.curselection()
		value = l1.get(int(l1.curselection()[0]))
		old_csv = read_csv_full('datas/query_history.csv', delimiter='&')
		new_csv = []
		for row in old_csv:
			if value not in row:
				new_csv.append(row)
			else:
				answer = messagebox.askokcancel('Предупреждение', 'Вы точно хотите удалить?', parent=history_window_0)
				if answer == False:
					new_csv.append(row)
				else:
					t0['state'] = tk.NORMAL
					t0.delete(0.0, tk.END)
					l1.delete(selection[0])
					t0['state'] = tk.DISABLED
		write_history(new_csv, type_record='w')

		sel = l0.curselection()
		val = l0.get(int(l0.curselection()[0]))
		dict_for_history_set = make_dict_for_history_set()[0]
		if dict_for_history_set[val] == []:
			l0.delete(sel[0])

	def choose_code_l1(evt):
		if t0.grid_info() == {}:
			t0.grid(row=1, column=3, columnspan=2, stick='ws')
		t0['state'] = tk.NORMAL
		t0.delete(0.0, tk.END)
		w = evt.widget
		value = w.get(int(w.curselection()[0]))
		for i, row in enumerate(history_dict[value]):
			t0.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n')
		t0['state'] = tk.DISABLED
		b1.grid(row=2, column=2, stick='w')
		b2.grid(row=2, column=3, stick='w')
		b3.grid(row=2, column=4, stick='e')

	def show_codes_l0(evt):
		sample_codes_code = []
		w = evt.widget
		value = w.get(int(w.curselection()[0]))
		dict_for_history_set = make_dict_for_history_set()[0]
		for i, row in enumerate(dict_for_history_set[value]):
			sample_codes_code.append(row)
		list_var_code_from_codes = tk.Variable(value=sample_codes_code)
		l1['listvariable'] = list_var_code_from_codes
		l1.grid(row=1, column=2, stick='ws')
		l1.selection_clear(0, tk.END)
		b1.grid_forget()
		b2.grid_forget()
		b3.grid_forget()
		t0.grid_forget()

	def cb_search(event=None):
		sstr = search_str.get()
		l0.delete(0, tk.END)
		# If filter removed show all data
		if sstr == "":
			fill_listbox(sample_codes)
			return

		filtered_data = list()
		for item in sample_codes:
			if item.find(sstr) >= 0:
				filtered_data.append(item)

		fill_listbox(filtered_data)

	def cb_del():
		search_history.delete(0, tk.END)
		l0.delete(0, tk.END)
		fill_listbox(sample_codes)

	def fill_listbox(ld):
		for item in ld:
			l0.insert(tk.END, item)

	def make_dict_for_history_set():
		history_dict = dict_from_csv()
		history_samples = list(history_dict)[::-1]

		dict_for_history_set = defaultdict(list)
		for i in history_samples:
			try:
				code = i.split('-')
				if len(code) > 3:
					code = ('-').join(code[:-1])
				else:
					code = i
			except:
				code = i
			dict_for_history_set[code].append(i)
		for key in dict_for_history_set.keys():
			if len(dict_for_history_set[key]) > 1:
				dict_for_history_set[key] = sorted(dict_for_history_set[key], key=lambda y: int(y.split('-')[-1]))
		return dict_for_history_set, history_dict, history_samples

	infos_for_history = ['Номер лабораторного журнала', 'Регистрационный номер пробы', 'Наименование пробы(образца)',
	                     'ФИО специалиста ответственного за пробоподготовку', 'Примечания пробоподготовки',
	                     'Примечания регистрационного журнала', 'Перечень показателей через запятую',
	                     'Реквизиты НД для проведения пробоподготовки',
	                     'Реквизиты НД на метод исследования', 'ФИО специалиста проводившего исследование',
	                     'ФИО ответственного исполнителя', 'Дата начала исследования', 'Дата начала пробоподготовки',
	                     'Дата отбора пробы (образца)', 'Дата поступления', 'Дата окончания исследования',
	                     'Дата окончания пробоподготовки', 'Дата утилизации пробы/сведения о консервации',
	                     'Дата выписки листа протокола', 'Этапы пробоподготовки', 'Этапы исследования']

	dict_for_history_set, history_dict, history_samples = make_dict_for_history_set()
	sample_codes = [key for key in dict_for_history_set]

	history_listbox_frame = tk.Frame(history_window_0)

	search_str = tk.StringVar()
	search_history = tk.Entry(history_listbox_frame, textvariable=search_str, justify=tk.LEFT, width=20)
	search_history.grid(row=0, column=0, columnspan=7, stick='w', pady=5)
	search_history.bind('<Return>', cb_search)
	b_search_confirm = tk.Button(history_listbox_frame, text='Поиск', command=cb_search)
	b_search_confirm.grid(row=0, column=7, columnspan=2, stick='w', pady=5)
	b_search_delete = tk.Button(history_listbox_frame, text='Удалить', command=cb_del)
	b_search_delete.grid(row=0, column=9, columnspan=2, stick='w', pady=5)

	l0_scroll = tk.Scrollbar(history_listbox_frame, orient='vertical')
	# list_var_codes = tk.Variable(value=sample_codes)
	l0 = tk.Listbox(history_listbox_frame, height=25, font=('Calibri', '10'),
	                exportselection=False, yscrollcommand=l0_scroll.set) # exportselection отвечает за то, чтобы при работе с виджетом можно было работать с другим без вреда для первого и второг
	fill_listbox(sample_codes)
	l0_scroll.config(command=l0.yview)
	history_listbox_frame.grid(row=1, column=0, stick='w')
	l0_scroll.grid(row=1, column=0, stick='nsw')
	l0.grid(row=1, column=1, columnspan=10, stick='we')
	l0.bind('<<ListboxSelect>>', show_codes_l0)


	# exportselection отвечает за то, чтобы при работе с виджетом можно было работать с другим без вреда для первого и второго

	sample_codes_code = []
	list_var_code_from_codes = tk.Variable(value=sample_codes_code)
	l1 = tk.Listbox(history_window_0, height=25, font=('Calibri', '10'), listvariable=list_var_code_from_codes,
	                exportselection=False)  # exportselection отвечает за то, чтобы при работе с виджетом можно было работать с другим без вреда для первого и второго
	l1.bind('<<ListboxSelect>>', choose_code_l1)

	t0 = tk.Text(history_window_0, width=100, height=26, font=('Calibri', '10'), wrap=tk.WORD, state=tk.DISABLED)

	b0 = tk.Button(history_window_0, text='Сгенерировать ворд', font=('Arial', '10'), command=func_for_window)
	b0.grid(row=2, column=0, columnspan=2, stick='we')
	b1 = tk.Button(history_window_0, text='Удалить запись', font=('Arial', '10'), command=delete_from_csv)
	b2 = tk.Button(history_window_0, text='Заполнить только пустые', font=('Arial', '10'),
	               command=confirm_empty_to_main)
	b3 = tk.Button(history_window_0, text='Применить', font=('Arial', '10'), command=confirm_to_main)


def clear_all_information():
	for i in range(len(variables_for_row)):
		variables_for_row[i].delete(0, tk.END)


def clear_cell(index):
	variables_for_row[index].delete(0, tk.END)


def add_all_datas(load=True):
	book_1 = op.load_workbook(filename=base_path1)
	sheet_1 = book_1.active

	book_2 = op.load_workbook(filename=base_path2)
	sheet_2 = book_2.active

	raw_data_sample = []
	for i, row in enumerate(sheet_1.iter_rows(min_row=5, values_only=True)):
		formatted_row = []
		row = list(row)
		if str(row[0]).isdigit():
			row = row[0:15]
			for index in (5, 6, 10, 11):
				try:
					row[index] = row[index].strftime('%d.%m.%Y')
				except:
					pass
			for string in row:
				try:
					string = string.replace('\n', ' ')
				except (TypeError, AttributeError):
					pass
				try:
					formatted_row.append(string.strip())
				except:
					formatted_row.append(string)
			raw_data_sample.append(formatted_row)

	raw_data_register = []
	for i, row in enumerate(sheet_2.iter_rows(min_row=5, values_only=True)):
		formatted_row = []
		row = list(row)
		if str(row[0]).isdigit():
			row = row[0:12]
			for index in (3, 4, 5, 7, 8, 9):
				try:
					row[index] = row[index].strftime('%d.%m.%Y')
				except:
					pass
			for string in row:
				try:
					formatted_row.append(string.strip())
				except:
					formatted_row.append(string)
			raw_data_register.append(formatted_row)

	dict_sample = {}
	dict_register = {}
	for row in raw_data_sample:
		dict_key = 'неизвестный номер'
		if re.fullmatch(f"{r'^(.*)-(.*)-(.*)'}", row[1]):
			dict_key = row[1]
		else:
			for row_register in raw_data_register:
				if row[0] == row_register[0]:
					dict_key = row_register[1]
		dict_values = row
		dict_sample[dict_key] = dict_values
	for row in raw_data_register:
		dict_key = 'неизвестный номер'
		if re.fullmatch(f"{r'^(.*)-(.*)-(.*)'}", row[1]):
			dict_key = row[1]
		else:
			for row_sample in raw_data_sample:
				if row[0] == row_sample[0]:
					dict_key = row_sample[1]
		dict_values = row
		dict_register[dict_key] = dict_values

	sample_keys = set(dict_sample.keys())
	register_keys = set(dict_register.keys())

	full_dict = defaultdict(list)
	full_set = sample_keys | register_keys

	for code in full_set:
		try:
			full_dict[code].append(dict_sample[code])
		except KeyError:
			full_dict[code].append(['' for x in range(0, 15)])
		try:
			full_dict[code].append(dict_register[code])
		except KeyError:
			full_dict[code].append(['' for x in range(0, 12)])

	final_datas = []
	for key, value in full_dict.items():
		# 0 nb_lab_journal.get()
		# print(value)
		if value[0][0] == value[1][0]:
			cell0 = value[0][0]
		elif value[0][0] == '':
			cell0 = value[1][0]
		elif value[1][0] == '':
			cell0 = value[0][0]
		else:
			cell0 = value[0][0]

		# 1 rg_nb_sample.get() s[1] r[1]
		if value[0][1] == value[1][1]:
			cell1 = value[0][1]
		elif value[0][1] == '':
			cell1 = value[1][1]
		elif value[1][1] == '':
			cell1 = value[0][1]
		else:
			if re.fullmatch(f"{r'^(.*)-(.*)-(.*)'}", value[0][1]):
				cell1 = value[0][1]
			else:
				cell1 = value[1][1]

		# 2 name_sample.get() s[2] r[2]
		if value[0][2] == value[1][2]:
			cell2 = value[0][2]
		elif value[0][2] == '':
			cell2 = value[1][2]
		elif value[1][2] == '':
			cell2 = value[0][2]
		else:
			cell2 = value[0][2]

		# 3 nm_sample_executor.get() s[7]
		cell3 = value[0][7]

		# 4 nt_sample.get() s[14]
		cell4 = value[0][14]

		# 5 nt_register.get() r[11]
		cell5 = value[1][11]

		# 6 ls_indicators_research s[12]
		cell6 = value[0][12]

		# 7 det_nd_prep_sample.get() s[3]
		cell7 = value[0][3]

		# 8 det_nd_research.get() s[8]
		cell8 = value[0][8]

		# 9 sp_did_research.get() s[13]
		cell9 = value[0][13]

		# 10 rsp_executor.get() r[10]
		cell10 = value[1][10]

		# 11 dt_st_research.get() s[10] r[5]
		if value[0][10] == value[1][5]:
			cell11 = value[0][10]
		elif value[0][10] == '':
			cell11 = value[1][5]
		elif value[1][5] == '':
			cell11 = value[0][10]
		else:
			st_sample_date = value[0][10]
			st_register_date = value[1][5]
			try:
				datetime_st_sample_date = datetime.datetime.strptime(st_sample_date, '%d.%m.%Y')
			except:
				pass
			try:
				datetime_st_register_date = datetime.datetime.strptime(st_register_date, '%d.%m.%Y')
			except:
				pass
			if datetime_st_sample_date != '' and datetime_st_register_date != '':
				if datetime_st_sample_date > datetime_st_register_date:
					cell11 = st_sample_date
				else:
					cell11 = st_register_date
			elif datetime_st_sample_date == '' and datetime_st_register_date != '':
				cell11 = datetime_st_register_date
			elif datetime_st_sample_date != '' and datetime_st_register_date == '':
				cell11 = datetime_st_sample_date
			else:
				cell11 = '-'

		# 12 dt_st_sample_prep.get() s[5]
		cell12 = value[0][5]

		# 13 dt_st_sampling.get() r[3]
		cell13 = value[1][3]

		# 14 dt_get_receipt.get() r[4]
		cell14 = value[1][4]

		# 15 dt_fn_research.get() s[11] r[7]
		if value[0][11] == value[1][7]:
			cell15 = value[0][11]
		elif value[0][11] == '':
			cell15 = value[1][7]
		elif value[1][7] == '':
			cell15 = value[0][11]
		else:
			fn_sample_date = value[0][11]
			fn_register_date = value[1][7]
			datetime_fn_sample_date = ''
			datetime_fn_register_date = ''
			try:
				datetime_fn_sample_date = datetime.datetime.strptime(fn_sample_date, '%d.%m.%Y')
			except:
				pass
			try:
				datetime_fn_register_date = datetime.datetime.strptime(fn_register_date, '%d.%m.%Y')
			except:
				pass
			if datetime_fn_sample_date != '' and datetime_fn_register_date != '':
				if datetime_fn_sample_date > datetime_fn_register_date:
					cell15 = fn_sample_date
				else:
					cell15 = fn_register_date
			elif datetime_fn_sample_date == '' and datetime_fn_register_date != '':
				cell15 = datetime_fn_register_date
			elif datetime_fn_sample_date != '' and datetime_fn_register_date == '':
				cell15 = datetime_fn_sample_date
			else:
				cell15 = '-'

		# 16 dt_fn_sample_prep.get() s[6]
		cell16 = value[0][6]

		# 17 dt_disposal.get() r[8]
		cell17 = value[1][8]

		# 18 dt_issue_protocol.get() r[9]
		cell18 = value[1][9]

		# 19 steps_sample.get() s[4]
		cell19 = value[0][4]

		# 20 stp_research.get() s[9]
		cell20 = value[0][9]

		to_final_data = [
			cell0,  # nb_lab_journal.get(), s[0] r[0]
			cell1,  # rg_nb_sample.get(), s[1] r[1]
			cell2,  # name_sample.get(), s[2] r[2]
			cell3,  # nm_sample_executor.get(), s[7]
			cell4,  # nt_sample.get(), s[14]
			cell5,  # nt_register.get(), r[11]
			cell6,  # ls_indicators_research, s[12]
			cell7,  # det_nd_prep_sample.get(), s[3]
			cell8,  # det_nd_research.get(), s[8]
			cell9,  # sp_did_research.get(), s[13]
			cell10,  # rsp_executor.get(), r[10]
			cell11,  # dt_st_research.get(), s[10] r[5]
			cell12,  # dt_st_sample_prep.get(), s[5]
			cell13,  # dt_st_sampling.get(), r[3]
			cell14,  # dt_get_receipt.get(), r[4]
			cell15,  # dt_fn_research.get(), s[11] r[7]
			cell16,  # dt_fn_sample_prep.get(), s[6]
			cell17,  # dt_disposal.get(), r[8]
			cell18,  # dt_issue_protocol.get(), r[9]
			cell19,  # steps_sample.get(), s[4]
			cell20,  # stp_research.get(), s[9]
		]
		final_datas.append(to_final_data)

	final_datas.sort(key=lambda x: x[0])
	if load:
		with open('datas/query_history.csv', 'a', encoding='utf-8', newline='') as f:
			for row in final_datas:
				writer = csv.writer(f, delimiter='&')
				writer.writerow(row)
			f.close()
	if load == False:
		return final_datas


def refresh_changes():
	infos_for_history = ['Номер лабораторного журнала', 'Регистрационный номер пробы',
	                     'Наименование пробы(образца)',
	                     'ФИО специалиста ответственного за пробоподготовку', 'Примечания пробоподготовки',
	                     'Примечания регистрационного журнала', 'Перечень показателей через запятую',
	                     'Реквизиты НД для проведения пробоподготовки',
	                     'Реквизиты НД на метод исследования', 'ФИО специалиста проводившего исследование',
	                     'ФИО ответственного исполнителя', 'Дата начала исследования',
	                     'Дата начала пробоподготовки',
	                     'Дата отбора пробы (образца)', 'Дата поступления', 'Дата окончания исследования',
	                     'Дата окончания пробоподготовки', 'Дата утилизации пробы/сведения о консервации',
	                     'Дата выписки листа протокола', 'Этапы пробоподготовки', 'Этапы исследования']

	def choose_changes(evt):
		t0['state'] = tk.NORMAL
		t0.delete(0.0, tk.END)
		w = evt.widget
		value = w.get(int(w.curselection()[0]))

		try:
			for i, row in enumerate(query_history_dict_checker_from_csv[value]):
				if value in query_history_dict:
					if query_history_dict_checker_from_csv[value][i] != query_history_dict[value][i]:
						t0.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n', 'before')
					else:
						t0.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n')
				else:
					t0.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n', 'before')
		except KeyError:
			t0.insert(tk.INSERT, 'Запись будет добавлена в базу')
			status = 'добавление'
		t0['state'] = tk.DISABLED

		t1['state'] = tk.NORMAL
		t1.delete(0.0, tk.END)

		try:
			for i, row in enumerate(query_history_dict[value]):
				if value in query_history_dict_checker_from_csv:
					if query_history_dict[value][i] != query_history_dict_checker_from_csv[value][i]:
						t1.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n', 'after')
						status = 'изменения'
					else:
						t1.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n')
				else:
					t1.insert(tk.INSERT, infos_for_history[i] + ' - ' + row + '\n', 'after')
		except KeyError:
			t1.insert(tk.INSERT, 'Запись будет удалена из базы')
			status = 'удаление'
		t1['state'] = tk.DISABLED

	def too_much_changes():
		messagebox.showerror('Инфо',
		                     'Слишком много изменений для выбора. Воспользуйтесь кнопкой "Принять все изменения", чтобы принять все изменения.')

	def confirm_all_changes():
		write_history(query_history_dict.values(), type_record='w')
		messagebox.showinfo('Инфо', f'Было принято {counter_of_changes} изменений', parent=refresh_changes_window)

	def save_checkbox_to_csv():
		pass

	def choose_changes_checkbox_menu():
		window_for_choose_checkbox = tk.Toplevel(refresh_changes_window)  # нельзя нажимать в других окнах
		window_for_choose_checkbox.title('Окно 2 чекбоксы')
		window_for_choose_checkbox.geometry(f'{int(600.0 * scaling)}x{int(525.0 * scaling)}+1000+350')
		window_for_choose_checkbox.protocol('WM_DELETE_WINDOW')  # закрытие приложения

		def on_mousewheel(event):
			if isinstance(event.widget, tk.Checkbutton):
				scroll = -1 if event.delta > 0 else 1
				my_canvas.yview_scroll(scroll, 'units')

		def on_mousewheel_regular(event):
			scroll = -1 if event.delta > 0 else 1
			my_canvas.yview_scroll(scroll, 'units')

		# Create a main frame
		main_frame = tk.Frame(window_for_choose_checkbox)
		main_frame.pack(fill=tk.BOTH, expand=1)

		# Create a canvas
		my_canvas = tk.Canvas(main_frame)
		my_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=1)

		# Add a scrollbar to the canvas
		my_scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=my_canvas.yview)
		my_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

		# Configure the canvas
		my_canvas.configure(yscrollcommand=my_scrollbar.set)
		my_canvas.bind('<Configure>', lambda e: my_canvas.configure(scrollregion=my_canvas.bbox('all')))

		# Create ANOTHER frame INSIDE the canvas
		second_frame = tk.Frame(my_canvas)

		# Add that new frame to a window in the canvas
		my_canvas.create_window((0, 0), window=second_frame, anchor='nw')

		dict_for_changes_save = {}

		def confirm_changes_checkbox_menu():
			write_history(query_history_dict.values(), type_record='w')

		def set_checkbox_on(index):
			if checkbox_variable_dict['checkbox_code_var_' + str(index)].get() == 'No':
				dict_for_changes_save[query_history_changes[index]] = 'No'
			if checkbox_variable_dict['checkbox_code_var_' + str(index)].get() == 'Yes':
				dict_for_changes_save[query_history_changes[index]] = 'Yes'

		def save_checkboxes_to_base():
			for key in dict_for_changes_save.keys():
				if dict_for_changes_save[key] == 'No':
					try:
						query_history_dict[key] = query_history_dict_checker_from_csv[key]
					except KeyError:
						del query_history_dict[key]
			counter_of_changes_to_save = [x for x in dict_for_changes_save.values() if x == 'Yes']
			askorcancel_save_checkboxes_to_base = messagebox.askokcancel('title',
			                                                             'Вы действительно хотите принять измения?',
			                                                             parent=window_for_choose_checkbox)
			if askorcancel_save_checkboxes_to_base == True:
				write_history(query_history_dict.values(), type_record='w')
				messagebox.showinfo('Инфо', f'Принято для {len(counter_of_changes_to_save)} изменений',
				                    parent=window_for_choose_checkbox)
				window_for_choose_checkbox.destroy()
				refresh_changes_window.destroy()

		def func_for_pass():
			pass

		index_for_column = 0
		index_for_row = 1
		checkbox_variable_dict = {}
		x = len(query_history_changes)
		if x <= 150:
			my_canvas.configure(yscrollcommand=None)
		else:
			my_canvas.bind_all('<MouseWheel>', on_mousewheel)
			my_canvas.bind('<MouseWheel>', on_mousewheel_regular)
		x_rows = 25
		if x > 150:
			if x % 6 == 0:
				x_rows = int(x / 6)
			else:
				x_rows = int(x / 6) + 1
		for i, code in enumerate(query_history_changes):
			if i != 0 and i % x_rows == 0:
				index_for_row = 1
				index_for_column += 1
			checkbox_variable_dict['checkbox_code_var_' + str(i)] = tk.StringVar()
			checkbox_variable_dict['checkbox_code_var_' + str(i)].set('Yes')
			dict_for_changes_save[query_history_changes[i]] = 'Yes'
			checkbox_variable_dict['checkbox_code' + str(i)] = tk.Checkbutton(
				second_frame, text=f'{code}', variable=checkbox_variable_dict['checkbox_code_var_' + str(i)],
				command=lambda index=i: set_checkbox_on(index), offvalue='No', onvalue='Yes')
			checkbox_variable_dict['checkbox_code' + str(i)].grid(row=index_for_row, column=index_for_column,
			                                                      stick='w')
			index_for_row += 1

		def uncheckbox_all():
			if checkboxes_on_off['text'] == 'Снять все галочки':
				for i, key in enumerate(dict_for_changes_save):
					checkbox_variable_dict['checkbox_code_var_' + str(i)].set('No')
					dict_for_changes_save[key] = 'No'
				checkboxes_on_off['text'] = 'Выделить все галчоки'
			elif checkboxes_on_off['text'] == 'Выделить все галчоки':
				for i, key in enumerate(dict_for_changes_save):
					checkbox_variable_dict['checkbox_code_var_' + str(i)].set('Yes')
					dict_for_changes_save[key] = 'Yes'
				checkboxes_on_off['text'] = 'Снять все галочки'

		tk.Button(second_frame, text='Сохранить в базу', command=save_checkboxes_to_base).grid(
			row=0, column=0, stick='w')
		checkboxes_on_off = tk.Button(second_frame, text='Снять все галочки', command=uncheckbox_all)
		checkboxes_on_off.grid(row=0, column=2, stick='w')

	def close_refresh_changes_window():
		refresh_changes_window.destroy()

	query_history_dict = dict_from_csv()
	query_history_dict_checker_from_csv = query_history_dict.copy()

	dict_from_excel = {}
	for row in add_all_datas(load=False):
		dict_key = row[1]
		row[0] = str(row[0])
		dict_values = ['' if v is None else v for v in row]
		dict_from_excel[dict_key] = dict_values

	dict_from_excel_codes = dict_from_excel.keys()
	counter_of_changes = 0
	query_history_changes = []
	for code in dict_from_excel_codes:
		try:
			str_1 = ('&').join(dict_from_excel[code])
			try:
				str_2 = ('&').join(query_history_dict[code])
			except KeyError:
				str_2 = None
			if str_1 != str_2:
				query_history_dict[code] = dict_from_excel[code]
				query_history_changes.append(code)
				counter_of_changes += 1
		except TypeError:
			pass

	query_history_dict_keys = list(query_history_dict.keys())
	for key in query_history_dict_keys:
		if key in dict_from_excel_codes:
			pass
		else:
			del query_history_dict[key]
			query_history_changes.append(key)
			counter_of_changes += 1
	if counter_of_changes == 0:
		messagebox.showinfo('готово', 'Изменений нет', parent=win)
	else:
		refresh_changes_window = tk.Toplevel(win)
		refresh_changes_window.title('1 меню изменений')
		refresh_changes_window.geometry(f'{int(1225.0 * scaling)}x{int(500.0 * scaling)}+1200+400')
		refresh_changes_window.protocol('WM_DELETE_WINDOW')  # закрытие приложения

		changes_frame = tk.Frame(refresh_changes_window)
		l0_scroll = tk.Scrollbar(changes_frame, orient='vertical')
		list_var_changes = tk.Variable(value=query_history_changes)
		l0 = tk.Listbox(changes_frame, height=19, listvariable=list_var_changes,
		                exportselection=False, yscrollcommand=l0_scroll.set)  # exportselection отвечает за то, чтобы при работе с виджетом можно было работать с другим без вреда для первого и второго

		l0_scroll.config(command=l0.yview)
		changes_frame.grid(row=0, column=0, stick='e')
		l0_scroll.grid(row=0, column=0, stick='ens')
		l0.grid(row=0, column=1, stick='w')
		l0.bind('<<ListboxSelect>>', choose_changes)

		t0 = tk.Text(refresh_changes_window, width=100, wrap=tk.WORD, state=tk.DISABLED)
		t0.tag_configure("before", foreground="red", background='#FFFFDA')
		t0.grid(row=0, column=2, padx=5)

		t1 = tk.Text(refresh_changes_window, width=100, wrap=tk.WORD, state=tk.DISABLED)
		t1.tag_configure("after", foreground="green", background='#FFFFDA')
		t1.grid(row=0, column=3, padx=5)

		tk.Button(refresh_changes_window, text='Отменить изменения', command=close_refresh_changes_window).grid(
			row=1, column=0, columnspan=2, stick='we')

		if len(query_history_changes) <= 600:
			tk.Button(refresh_changes_window, text='Выбрать изменения', command=choose_changes_checkbox_menu).grid(
				row=1, column=2, stick='w', padx=5)
		else:
			tk.Button(refresh_changes_window, text='Выбрать изменения', command=too_much_changes).grid(
				row=1, column=2, stick='w', padx=5)

		tk.Button(refresh_changes_window, text='Принять все изменения', command=confirm_all_changes).grid(
			row=1, column=3, stick='w', padx=5)


# двойное
tk.Label(win, text='Номер лабораторного журнала').grid(row=0, column=0, stick='e')
nb_lab_journal = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
nb_lab_journal.grid(row=0, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(0),
          borderwidth=0).grid(
	row=0, column=2, stick='w', padx=5)
tk.Button(win, text='Очистить все', command=clear_all_information).grid(row=0, column=3, stick='w')
tk.Button(win, text='История', command=history_window).grid(row=0, column=4, stick='w')
tk.Button(win, text='Настройки', command=settings_window).grid(row=0, column=5, stick='w')

# двойное
tk.Label(win, text='Регистрационный номер пробы').grid(row=1, column=0, stick='e')
rg_nb_sample = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
rg_nb_sample.grid(row=1, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(1),
          borderwidth=0).grid(
	row=1, column=2, stick='w', padx=5)

# двойное
tk.Label(win, text='Наименование пробы(образца)').grid(row=2, column=0, stick='e')
name_sample = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
name_sample.grid(row=2, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(2),
          borderwidth=0).grid(
	row=2, column=2, stick='w', padx=5)

# уникальное
tk.Label(win, text='ФИО специалиста ответственного за пробоподготовку').grid(row=3, column=0, stick='e')
nm_sample_executor = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
nm_sample_executor.grid(row=3, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(3),
          borderwidth=0).grid(
	row=3, column=2, stick='w', padx=5)
tk.Button(win, text='Выбрать специалиста',
          command=lambda: start_window_0(nm_sample_executor, 'datas/nm_sample_executor.csv')).grid(row=3, column=3,
                                                                                                   stick='w')

# уникальное
tk.Label(win, text='Примечания пробоподготовки').grid(row=4, column=0, stick='e')
nt_sample = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
nt_sample.grid(row=4, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(4),
          borderwidth=0).grid(
	row=4, column=2, stick='w', padx=5)

# уникальное
tk.Label(win, text='Примечания регистрационного журнала').grid(row=5, column=0, stick='e')
nt_register = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
nt_register.grid(row=5, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(5),
          borderwidth=0).grid(
	row=5, column=2, stick='w', padx=5)

# Перечень показателей
default_indicator = 'не обнаружены'
list_of_indicators = ('не обнаружены', 'обнаружены')
tk.Label(win, text='Перечень показателей через запятую').grid(row=6, column=0, stick='e')
ls_indicators = tk.Entry(win, font=('Arial', 10), width=25)
ls_indicators.grid(row=6, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(6),
          borderwidth=0).grid(
	row=6, column=2, stick='w', padx=5)
combo_indicators = ttk.Combobox(win, values=list_of_indicators)
combo_indicators.current(0)
combo_indicators.grid(row=6, column=3, stick='w')
combo_indicators.bind("<<ComboboxSelected>>", find_not_find)

# Реквизиты НД для проведения пробоподготовки и на метод исследования
repeat_for_nd_value = tk.StringVar()
repeat_for_nd_value.set('No')
tk.Label(win, text='Реквизиты НД для проведения пробоподготовки').grid(row=7, column=0, stick='e')
det_nd_prep_sample = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
det_nd_prep_sample.grid(row=7, column=1)
det_nd_prep_sample.bind("<FocusOut>", nd_check_button_off)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(7),
          borderwidth=0).grid(
	row=7, column=2, stick='w', padx=5)
tk.Label(win, text='Реквизиты НД на метод исследования').grid(row=8, column=0, stick='e')
det_nd_research = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
det_nd_research.bind("<FocusIn>", nd_check_button_off)
det_nd_research.bind("<FocusOut>", nd_check_button_off)
det_nd_research.grid(row=8, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(8),
          borderwidth=0).grid(
	row=8, column=2, stick='w', padx=5)
nd_check_button = tk.Checkbutton(win, text='повторить реквизиты НД пробоподготовки', command=repeat_for_nd,
                                 variable=repeat_for_nd_value, offvalue='No', onvalue='Yes')
nd_check_button.grid(row=8, column=3, stick='w')

# ФИО специалиста и ответственный исполнитель
repeat_for_sp_value = tk.StringVar()
repeat_for_sp_value.set('No')
tk.Label(win, text='ФИО специалиста проводившего исследование').grid(row=9, column=0, stick='e')
sp_did_research = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
sp_did_research.grid(row=9, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(9),
          borderwidth=0).grid(
	row=9, column=2, stick='w', padx=5)
tk.Button(win, text='Выбрать специалиста',
          command=lambda: start_window_0(sp_did_research, 'datas/sp_did_research.csv')).grid(row=9, column=3,
                                                                                             stick='w')

tk.Label(win, text='ФИО ответственного исполнителя').grid(row=10, column=0, stick='e')
rsp_executor = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
rsp_executor.grid(row=10, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(10),
          borderwidth=0).grid(row=10, column=2, stick='w', padx=5)
tk.Button(win, text='Выбрать специалиста',
          command=lambda: start_window_0(rsp_executor, 'datas/rsp_executor.csv')).grid(
	row=10, column=3, stick='w')

# Даты начала

# двойное
tk.Label(win, text='Дата начала исследования').grid(row=11, column=0, stick='e')
dt_st_research = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_st_research.grid(row=11, column=1)
dt_st_research.bind("<FocusOut>", check_st_functions)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(11),
          borderwidth=0).grid(row=11, column=2, stick='w', padx=5)

dt_st_value_1 = tk.StringVar()
dt_st_value_2 = tk.StringVar()
dt_st_value_3 = tk.StringVar()
dt_st_value_1.set('No')
dt_st_value_2.set('No')
dt_st_value_3.set('No')

tk.Label(win, text='Дата начала пробоподготовки').grid(row=12, column=0, stick='e')
dt_st_sample_prep = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_st_sample_prep.grid(row=12, column=1)
dt_st_sample_prep.bind("<FocusIn>", dt_st_1_check_off)
dt_st_sample_prep.bind("<FocusOut>", dt_st_1_check_off)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(12),
          borderwidth=0).grid(row=12, column=2, stick='w', padx=5)
dt_st_check_button_1 = tk.Checkbutton(win, text='повторить дату начала исследования', command=repeat_for_dt_st_1,
                                      variable=dt_st_value_1, offvalue='No', onvalue='Yes')
dt_st_check_button_1.grid(row=12, column=3, stick='w')

tk.Label(win, text='Дата отбора пробы (образца)').grid(row=13, column=0, stick='e')
dt_st_sampling = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_st_sampling.grid(row=13, column=1)
dt_st_sampling.bind("<FocusIn>", dt_st_2_check_off)
dt_st_sampling.bind("<FocusOut>", dt_st_2_check_off)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(13),
          borderwidth=0).grid(row=13, column=2, stick='w', padx=5)
dt_st_check_button_2 = tk.Checkbutton(win, text='повторить дату начала исследования', command=repeat_for_dt_st_2,
                                      variable=dt_st_value_2, offvalue='No', onvalue='Yes')
dt_st_check_button_2.grid(row=13, column=3, stick='w')

tk.Label(win, text='Дата поступления').grid(row=14, column=0, stick='e')
dt_get_receipt = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_get_receipt.grid(row=14, column=1)
dt_get_receipt.bind("<FocusIn>", dt_st_3_check_off)
dt_get_receipt.bind("<FocusOut>", dt_st_3_check_off)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(14),
          borderwidth=0).grid(row=14, column=2, stick='w', padx=5)
dt_st_check_button_3 = tk.Checkbutton(win, text='повторить дату начала исследования', command=repeat_for_dt_st_3,
                                      variable=dt_st_value_3, offvalue='No', onvalue='Yes')
dt_st_check_button_3.grid(row=14, column=3, stick='w')

# Даты окончания

tk.Label(win, text='Дата окончания исследования').grid(row=15, column=0, stick='e')
dt_fn_research = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_fn_research.grid(row=15, column=1)
dt_fn_research.bind("<FocusOut>", check_fn_functions)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(15),
          borderwidth=0).grid(row=15, column=2, stick='w', padx=5)

dt_fn_value_1 = tk.StringVar()
dt_fn_value_2 = tk.StringVar()
dt_fn_value_3 = tk.StringVar()
dt_fn_value_1.set('No')
dt_fn_value_2.set('No')
dt_fn_value_3.set('No')

tk.Label(win, text='Дата окончания пробоподготовки').grid(row=16, column=0, stick='e')
dt_fn_sample_prep = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_fn_sample_prep.grid(row=16, column=1)
dt_fn_sample_prep.bind("<FocusIn>", dt_fn_1_check_off)
dt_fn_sample_prep.bind("<FocusOut>", dt_fn_1_check_off)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(16),
          borderwidth=0).grid(row=16, column=2, stick='w', padx=5)
dt_fn_check_button_1 = tk.Checkbutton(win, text='повторить дату окончания исследования', command=repeat_for_dt_fn_1,
                                      variable=dt_fn_value_1, offvalue='No', onvalue='Yes')
dt_fn_check_button_1.grid(row=16, column=3, stick='w')

tk.Label(win, text='Дата утилизации пробы/сведения о консервации').grid(row=17, column=0, stick='e')
dt_disposal = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_disposal.grid(row=17, column=1)
dt_disposal.bind("<FocusIn>", dt_fn_2_check_off)
dt_disposal.bind("<FocusOut>", dt_fn_2_check_off)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(17),
          borderwidth=0).grid(row=17, column=2, stick='w', padx=5)
dt_fn_check_button_2 = tk.Checkbutton(win, text='повторить дату окончания исследования',
                                      command=repeat_for_dt_fn_2, variable=dt_fn_value_2, offvalue='No',
                                      onvalue='Yes')
dt_fn_check_button_2.grid(row=17, column=3, stick='w')

tk.Label(win, text='Дата выписки листа протокола').grid(row=18, column=0, stick='e')
dt_issue_protocol = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
dt_issue_protocol.grid(row=18, column=1)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(18),
          borderwidth=0).grid(row=18, column=2, stick='w', padx=5)
dt_fn_check_button_3 = tk.Checkbutton(win, text='повторить дату окончания исследования', command=repeat_for_dt_fn_3,
                                      variable=dt_fn_value_3, offvalue='No', onvalue='Yes')
dt_fn_check_button_3.grid(row=18, column=3, stick='w')
dt_issue_protocol.bind("<FocusIn>", dt_fn_3_check_off)
dt_issue_protocol.bind("<FocusOut>", dt_fn_3_check_off)
# Этапы исследования
repeat_for_stp_value = tk.StringVar()
repeat_for_stp_value.set('No')
tk.Label(win, text='Этапы пробоподготовки').grid(row=19, column=0, stick='e')
steps_sample = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
steps_sample.grid(row=19, column=1)
steps_sample.bind("<FocusOut>", check_stp_function)
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(19),
          borderwidth=0).grid(row=19, column=2, stick='w', padx=5)
tk.Label(win, text='Этапы исследования').grid(row=20, column=0, stick='e')
tk.Button(win, text='x', activeforeground='red', foreground='black', command=lambda: clear_cell(20),
          borderwidth=0).grid(row=20, column=2, stick='w', padx=5)
stp_research = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25)
stp_research.grid(row=20, column=1)
stp_check_button = tk.Checkbutton(win, text='заполнить этапы исследования', command=repeat_for_stp,
                                  variable=repeat_for_stp_value, offvalue='No', onvalue='Yes')
glb_stp_research_check_name = ''  # глобальная переменная для сверки и работы с галочкой
glb_stp_number_of_research = ''
stp_check_button.grid(row=20, column=3, stick='w')
stp_research.bind("<FocusIn>", for_stp_check_off)
stp_research.bind("<FocusOut>", for_stp_check_off)
# Эксель файл 1
tk.Label(win, text='Выбери эксель пробоподготовки 1').grid(row=21, column=0, stick='e')
e1_path = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25, state=tk.DISABLED)
tk.Button(text='Выберите файл', bd=5, font=('Arial', 10), command=get_file_1).grid(row=21, column=2, columnspan=2,
                                                                                   stick='w', padx=3)
e1_path.grid(row=21, column=1, stick='w')
# Эксель файл 2
tk.Label(win, text='Выбери регистрационный эксель файл 2').grid(row=22, column=0, stick='e')
e2_path = tk.Entry(win, justify=tk.LEFT, font=('Arial', 10), width=25, state=tk.DISABLED)
e2_path.grid(row=22, column=1, stick='w')
tk.Button(text='Выберите файл 2', bd=5, font=('Arial', 10), command=get_file_2).grid(row=22, column=2, columnspan=2,
                                                                                     stick='w', padx=3)
op_xl_button_value = tk.StringVar()
op_xl_button_value.set('No')
op_xl_button = tk.Checkbutton(win, text='открыть эксель', variable=op_xl_button_value, offvalue='No', onvalue='Yes')
op_xl_button.grid(row=22, column=4, stick='w')

# Кнопка на сервер
tk.Button(text='Добавить в excel', bd=5, font=('Arial', 10), command=excel_func).grid(row=100, column=0, stick='e',
                                                                                      pady=10)
# tk.Button(text='Сгенерировать word файл', bd=5, font=('Arial', 10), command=start_window_for_word).grid(
# 	row=100, column=1, stick='e', pady=10)
tk.Button(win, text='Добавить в историю из excel', command=refresh_changes).grid(row=100, column=4,
                                                                                 stick='w')

variables_for_row = [nb_lab_journal, rg_nb_sample, name_sample, nm_sample_executor, nt_sample, nt_register,
                     ls_indicators, det_nd_prep_sample, det_nd_research, sp_did_research, rsp_executor,
                     dt_st_research, dt_st_sample_prep, dt_st_sampling, dt_get_receipt, dt_fn_research,
                     dt_fn_sample_prep, dt_disposal, dt_issue_protocol, steps_sample, stp_research, ]

win.mainloop()
